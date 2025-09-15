import re

from jira import JIRA
from configparser import ConfigParser
import pandas as pd
from datetime import datetime, timedelta
import argparse
import codecs
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook
from team_performance import calculate_team_performance, export_team_performance_to_excel, add_team_performance_to_docx

# Configuration constants for accessing config.ini
CONFIG_FILE = "config.ini"
CONFIG_SECTION = "jira"
CONFIG_URL = "jira-url"
CONFIG_USERNAME = "username"
CONFIG_PASSWORD = "password"


def parse_arguments_and_config():
    """
    Parse command-line arguments and configuration file to retrieve JIRA credentials and project details.

    Returns:
        tuple: (jira_url, jira_username, jira_password, project, start_date, end_date, include_empty_weeks, member_list_file)
    """
    parser = argparse.ArgumentParser(description="Generate Jira report with custom date range.")
    parser.add_argument("-c", "--config", default=CONFIG_FILE, help="Path to config file (default: config.ini).")
    parser.add_argument("-u", "--username", help="Jira username (overrides config).")
    parser.add_argument("-p", "--password", help="Jira password (overrides config).")
    parser.add_argument("-l", "--url", help="Jira base URL (overrides config).")
    parser.add_argument("-proj", "--project", required=True, help="Jira project key (e.g., ABC).")
    parser.add_argument("--start-date", required=True, help="Start date in YYYY-MM-DD format.")
    parser.add_argument("--end-date", required=True, help="End date in YYYY-MM-DD format.")
    parser.add_argument("--include-empty-weeks", type=bool, default=True, help="Include weeks with no activity for all assignees.")
    parser.add_argument("--member-list-file", help="Path to Excel file with team member details.")
    parser.add_argument("--pr-stat-file", help="Path to Excel file with PR statistics.")
    args = parser.parse_args()

    # Load credentials from config file if not provided via arguments
    config = ConfigParser(allow_no_value=False, comment_prefixes=('#', ';'), inline_comment_prefixes='#')
    config.read_file(codecs.open(args.config, 'r', encoding='utf-8-sig'))

    jira_url = args.url or config.get(CONFIG_SECTION, CONFIG_URL, fallback=None)
    jira_username = args.username or config.get(CONFIG_SECTION, CONFIG_USERNAME, fallback=None)
    jira_password = args.password or config.get(CONFIG_SECTION, CONFIG_PASSWORD, fallback=None)

    # Validate required credentials
    if not jira_url or not jira_username or not jira_password:
        raise ValueError("Jira URL, username, and password must be specified in arguments or config file.")

    return jira_url, jira_username, jira_password, args.project, args.start_date, args.end_date, args.include_empty_weeks, args.member_list_file, args.pr_stat_file


def read_member_list(member_list_file):
    """
    Read team member details from an Excel file.

    Args:
        member_list_file (str): Path to Excel file with member data.

    Returns:
        list: Unique list of assignee names from column 'E', starting from row 2.

    Requirements:
        - Excel file must have a column 'E' with assignee names.
    """
    wb = load_workbook(member_list_file)
    sheet = wb.active
    assignee_column = 'E'

    # Extract unique assignees from column 'E', skipping header
    assignees = [sheet[f"{assignee_column}{row}"].value for row in range(2, sheet.max_row + 1) if sheet[f"{assignee_column}{row}"].value]
    return list(set(assignees))


def get_all_worklogs(jira, issue_key):
    """
    Fetch all worklogs for an issue using pagination.

    Args:
        jira (JIRA): Jira API client instance.
        issue_key (str): Jira issue key (e.g., JIRA-123).

    Returns:
        list: List of worklog entries for the issue.
    """
    worklogs = []
    start_at = 0
    while True:
        response = jira._session.get(
            f"{jira._options['server']}/rest/api/2/issue/{issue_key}/worklog",
            params={"startAt": start_at, "maxResults": 100}
        )
        response.raise_for_status()
        data = response.json()
        worklogs.extend(data.get("worklogs", []))
        if len(worklogs) >= data.get("total", 0):
            break
        start_at += 100
    return worklogs

def fetch_jira_data(jira, project, start_date, end_date):
    """
    Fetch Jira issues and worklogs for a project within a date range.

    Args:
        jira (JIRA): Jira API client instance.
        project (str): Jira project key.
        start_date (str): Start date in YYYY-MM-DD format.
        end_date (str): End date in YYYY-MM-DD format.

    Returns:
        pd.DataFrame: DataFrame with issue details, including key, summary, assignee, status,
                      resolution date, week, epic link, parent, type, description, labels,
                      priority, creator, and links.

    Requirements:
        - Jira API must provide fields: key, summary, assignee, resolutiondate, updated,
          customfield_10000 (epic link), parent, issuetype, description, labels, priority,
          creator, issuelinks.
        - Dates in resolutiondate and updated fields should be in YYYY-MM-DD format.
    """
    start_date = datetime.strptime(start_date, "%Y-%m-%d")
    end_date = datetime.strptime(end_date, "%Y-%m-%d")

    # Pagination variables
    start_at = 0
    max_results = 100
    all_issues = []

    # Fetch issues updated in the date range with pagination
    while True:
        jql_query = (
            f"project = {project} AND updated >= '{start_date.strftime('%Y-%m-%d')}' AND resolution in (Done, Resolved, Unresolved)"
        )
        issues = jira.search_issues(jql_query, startAt=start_at, maxResults=max_results, fields=[
            "key", "summary", "assignee", "resolutiondate", "updated", "customfield_10000",
            "parent", "issuetype", "description", "labels", "priority", "creator", "issuelinks"
        ])
        all_issues.extend(issues)
        if len(issues) < max_results:
            break
        start_at += max_results

    # Fetch epic names for epic links
    epic_keys = list({getattr(issue.fields, "customfield_10000", None) for issue in all_issues if getattr(issue.fields, "customfield_10000", None)})
    epic_names = {}
    if epic_keys:
        epics = jira.search_issues(f"issuekey in ({', '.join(epic_keys)})", maxResults=1000, fields=["key", "summary"])
        epic_names = {epic.key: epic.fields.summary for epic in epics}

    data = []
    for issue in all_issues:
        # Extract basic issue details
        key = issue.key
        summary = issue.fields.summary
        assignee = issue.fields.assignee.displayName if issue.fields.assignee else "Unassigned"
        resolved_date = issue.fields.resolutiondate
        epic_link = getattr(issue.fields, "customfield_10000", None)
        parent = getattr(issue.fields, "parent", None)
        parent_key = parent.key if parent else None
        parent_summary = parent.fields.summary if parent else None
        issue_type = getattr(issue.fields, "issuetype", None)
        issue_type_name = issue_type.name if issue_type else "Unknown"
        description = issue.fields.description or ""
        labels = ",".join(issue.fields.labels) if issue.fields.labels else ""
        priority = issue.fields.priority.name if issue.fields.priority else ""
        creator = issue.fields.creator.displayName if issue.fields.creator else ""

        # Process issue links (e.g., "relates to:JIRA-123,blocks:JIRA-456")
        links = ",".join(
            [f"{link.get('type', {}).get('name', '')}:{link.get('outwardIssue', {}).get('key', '')}"
             for link in issue.fields.issuelinks if "outwardIssue" in link] +
            [f"{link.get('type', {}).get('name', '')}:{link.get('inwardIssue', {}).get('key', '')}"
             for link in issue.fields.issuelinks if "inwardIssue" in link]
        ) if hasattr(issue.fields, "issuelinks") else ""

        # Fetch worklogs for the issue
        worklogs = get_all_worklogs(jira, key)
        worklog_dates = set()
        for log in worklogs:
            try:
                log_date = datetime.strptime(log["started"].split("T")[0], "%Y-%m-%d")
                if start_date <= log_date <= end_date:
                    worklog_dates.add(log_date)
            except Exception:
                continue

        resolved_week = None
        if resolved_date:
            resolution_date = resolved_date.split("T")[0]
            resolved_date_dt = datetime.strptime(resolution_date, "%Y-%m-%d")
            if start_date <= resolved_date_dt <= end_date:
                resolved_week = resolved_date_dt.strftime("%G-W%V")
                data.append({
                    "Issue_key": key,
                    "Summary": summary,
                    "Assignee": assignee,
                    "Status": "Resolved",
                    "Resolution_Date": resolution_date,
                    "Week": resolved_week,
                    "Epic_Link": epic_link,
                    "Epic_Name": epic_names.get(epic_link, "Unknown Epic"),
                    "Parent_Key": parent_key,
                    "Parent_Summary": parent_summary,
                    "Type": issue_type_name,
                    "Description": description,
                    "Labels": labels,
                    "Priority": priority,
                    "Creator": creator,
                    "Links": links
                })

        # Add worklog entries for in-progress tasks
        for log_date in worklog_dates:
            log_week = log_date.strftime("%G-W%V")
            if log_week != resolved_week:
                if not any(d["Issue_key"] == key and d["Week"] == log_week for d in data):
                    data.append({
                        "Issue_key": key,
                        "Summary": summary,
                        "Assignee": assignee,
                        "Status": "In progress",
                        "Week": log_week,
                        "Epic_Link": epic_link,
                        "Epic_Name": epic_names.get(epic_link, "Unknown Epic"),
                        "Parent_Key": parent_key,
                        "Parent_Summary": parent_summary,
                        "Type": issue_type_name,
                        "Description": description,
                        "Labels": labels,
                        "Priority": priority,
                        "Creator": creator,
                        "Links": links
                    })

    return pd.DataFrame(data)


def fill_missing_weeks(data, valid_weeks, required_assignees):
    """
    Add rows for assignees and weeks with no activity to ensure complete reporting.

    Args:
        data (pd.DataFrame): Jira data with Week and Assignee columns.
        valid_weeks (list): List of weeks in %G-W%V format.
        required_assignees (list): List of assignees to include.

    Returns:
        pd.DataFrame: DataFrame with added rows for missing weeks.
    """
    existing_keys = set(zip(data["Assignee"], data["Week"]))
    filler_rows = []

    for assignee in required_assignees:
        for week in valid_weeks:
            if (assignee, week) not in existing_keys:
                #year, week_num = map(int, week.split("-W"))
                filler_rows.append({
                    "Issue_key": "",
                    "Summary": "",
                    "Assignee": assignee,
                    "Status": "",
                    "Week": week,
                    "Epic_Link": "",
                    "Epic_Name": "",
                    "Parent_Key": "",
                    "Parent_Summary": "",
                    "Type": "",
                    "Description": "",
                    "Labels": "",
                    "Priority": "",
                    "Creator": "",
                    "Links": ""
                })

    if filler_rows:
        data = pd.concat([data, pd.DataFrame(filler_rows)], ignore_index=True)
    return data.sort_values(by=["Assignee", "Week"])

def generate_week_headers(valid_weeks, data):
    """
    Generate week headers for the report based on valid weeks and data.

    Args:
        valid_weeks (list): List of weeks in %G-W%V format.
        data (pd.DataFrame): Jira data with Week column.

    Returns:
        list: List of formatted week headers.
    """
    headers = []
    for week in valid_weeks:
        year, week_num = map(int, week.split("-W"))
        week_start = pd.Timestamp.fromisocalendar(year, week_num, 1).strftime("%Y-%m-%d")
        week_end = (pd.Timestamp.fromisocalendar(year, week_num, 1) + timedelta(days=6)).strftime("%Y-%m-%d")
        headers.append(f"ww{week_num} {week_start}–{week_end}")
    return headers

def generate_epic_report(data):
    """
    Generate a summary of resolved tasks grouped by epic.

    Args:
        data (pd.DataFrame): Jira data with Epic_Link, Epic_Name, Issue_key, Summary, Status.

    Returns:
        list: List of dictionaries with epic details and associated tasks.
    """
    epic_summary = []
    for epic_key, group in data[data["Status"] == "Resolved"].groupby("Epic_Link"):
        if epic_key:
            epic_summary.append({
                "Epic": group["Epic_Name"].iloc[0],
                "Tasks": [{"Task_Key": row["Issue_key"], "Task_Summary": row["Summary"]} for _, row in group.iterrows()]
            })
    return epic_summary

def add_hyperlink(paragraph, url, text, font_name="Calibri (Body)", font_size=10):
    """
    Add a hyperlink to a Word document paragraph.

    Args:
        paragraph: Word document paragraph object.
        url (str): URL for the hyperlink.
        text (str): Display text for the hyperlink.
        font_name (str): Font name for the hyperlink.
        font_size (int): Font size for the hyperlink.
    """
    part = paragraph.part
    r_id = part.relate_to(url, qn('r:hyperlink'), is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    text_run = OxmlElement('w:t')
    text_run.text = text
    new_run.append(text_run)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    run = paragraph.runs[-1]
    run.font.name = font_name
    run.font.size = Pt(font_size)

def set_paragraph_font(paragraph, font_name="Calibri (Body)", font_size=10):
    """
    Set font properties for a Word document paragraph.

    Args:
        paragraph: Word document paragraph object.
        font_name (str): Font name.
        font_size (int): Font size.
    """
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)

def generate_excel_report(data, start_date, end_date, project, headers, file_suffix):
    """
    Generate an Excel report with tasks organized by week and assignee.

    Args:
        data (pd.DataFrame): Jira data.
        start_date (datetime.date): Report start date.
        end_date (datetime.date): Report end date.
        project (str): Jira project key.
        headers (list): List of week headers.
        file_suffix (str): Suffix for output file name.
    """
    output_file = f"jira_report_{project}_{start_date}-{end_date}{file_suffix}.xlsx"
    data.to_excel(output_file, index=False)
    print(f"Excel report successfully created: {output_file}")

def add_resolved_tasks_section(document, resolved_tasks):
    """
    Add a section for resolved tasks to the Word document, grouped by week.

    Args:
        document: Word document object.
        resolved_tasks (pd.DataFrame): DataFrame with resolved tasks.
    """
    document.add_heading("Resolved Tasks", level=2)
    if resolved_tasks.empty:
        document.add_paragraph("No resolved tasks during the specified period.")
        return

    # Group tasks by week
    resolved_tasks = resolved_tasks.copy()  # Avoid chained assignment warnings
    resolved_tasks.loc[:, "Resolution_Week"] = pd.to_datetime(resolved_tasks["Resolution_Date"]).dt.strftime("%G-W%V")

    #resolved_tasks["Resolution_Week"] = pd.to_datetime(resolved_tasks["Resolution_Date"]).dt.strftime("%G-W%V")
    for week, tasks in resolved_tasks.groupby("Resolution_Week"):
        week_start = pd.to_datetime(f"{week}-1", format="%G-W%V-%u")
        week_end = week_start + timedelta(days=6)
        week_header = f"Week {week} ({week_start.strftime('%d/%m')} - {week_end.strftime('%d/%m')})"
        document.add_heading(week_header, level=3)

        for _, task in tasks[tasks["Type"] != "Sub-task"].iterrows():
            paragraph = document.add_paragraph(style="Normal")
            document.add_paragraph(f"{task['Issue_key']}: {task['Summary']}", style="List Bullet 2")
            set_paragraph_font(paragraph, font_name="Calibri (Body)", font_size=10)

            # List subtasks under parent task
            subtasks = tasks[tasks["Parent_Key"] == task["Issue_key"]]
            for _, subtask in subtasks.iterrows():
                document.add_paragraph(f"{subtask['Issue_key']}: {subtask['Summary']}", style="List Bullet 3")

def generate_word_report(data, start_date, end_date, project, headers, file_suffix, jira_url, epic_summary, member_list_file=None):
    """
    Generate a Word report with tabular view, list view, epic progress, and resolved tasks.

    Args:
        data (pd.DataFrame): Jira data.
        start_date (str): Report start date (YYYY-MM-DD).
        end_date (str): Report end date (YYYY-MM-DD).
        project (str): Jira project key.
        headers (list): List of week headers.
        file_suffix (str): Suffix for output file name.
        jira_url (str): Jira base URL for hyperlinks.
        epic_summary (list): Epic summary data.
        member_list_file (str, optional): Path to member list Excel file.

    Requirements:
        - Data must include columns: Issue_key, Summary, Assignee, Status, Week, Epic_Link, Epic_Name.
    """
    if member_list_file:
        required_assignees = read_member_list(member_list_file)
    else:
        required_assignees = data["Assignee"].unique()

    # Sort data for consistent reporting
    sorted_data = data.sort_values(by=["Assignee", "Week", "Status"])
    required_assignees.sort()

    document = Document()
    document.add_heading(f"JIRA Report: {project} - {start_date}-{end_date}", level=1)

    # Tabular View: Tasks by assignee and week
    document.add_heading("Tabular View", level=2)
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Add headers
    headers = ["Name", "Week #", "Date", "Description", "Link", "Status"]
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header

    # Add data rows, ensuring all required assignees are included
    for assignee in required_assignees:
        assignee_data = sorted_data[sorted_data["Assignee"] == assignee]

        if assignee_data.empty:
            # Add empty row for assignees with no tasks
            row_cells = table.add_row().cells
            row_cells[0].text = assignee
            row_cells[1].text = ""
            row_cells[2].text = ""
            row_cells[3].text = ""
            row_cells[4].text = ""
            row_cells[5].text = ""
        else:
            # Add rows for each task
            for _, row in assignee_data.iterrows():
                year, week_num = map(int, row["Week"].split("-W"))
                week_start = pd.Timestamp.fromisocalendar(year, week_num, 1).strftime("%Y/%m/%d")
                week_end = (pd.Timestamp.fromisocalendar(year, week_num, 1) + timedelta(days=6)).strftime("%Y/%m/%d")
                week_range = f"{week_start}–{week_end}"
                row_cells = table.add_row().cells
                row_cells[0].text = assignee
                row_cells[1].text = row["Week"]
                row_cells[2].text = week_range
                row_cells[3].text = row["Summary"]
                add_hyperlink(row_cells[4].paragraphs[0], f"{jira_url}/browse/{row['Issue_key']}", f"{row['Issue_key']}", font_size=8)
                row_cells[5].text = row["Status"]
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_font(paragraph, font_name="Calibri (Body)", font_size=8)

    # List View: Resolved tasks by assignee and week
    document.add_heading("List View", level=2)
    resolved_data = data[data["Status"] == "Resolved"]
    for assignee, group in resolved_data.groupby("Assignee"):
        paragraph_assignee = document.add_paragraph(assignee, style="Heading 2")
        set_paragraph_font(paragraph_assignee, font_name="Times New Roman", font_size=11)

        for week, week_data in group.groupby("Week"):
            year, week_num = map(int, week.split("-W"))
            week_start = pd.Timestamp.fromisocalendar(year, week_num, 1).strftime("%Y-%m-%d")
            week_end = (pd.Timestamp.fromisocalendar(year, week_num, 1) + timedelta(days=6)).strftime("%Y-%m-%d")
            week_header = f"ww{int(re.search(r'W(\d+)', week).group(1))} {week_start}-{week_end}"
            paragraph = document.add_paragraph(week_header, style="Heading 3")
            set_paragraph_font(paragraph, font_name="Times New Roman", font_size=11)

            # Add tasks for the current week
            for idx, row in enumerate(week_data.itertuples(index=False, name="Row"), start=1):
                # Add a new paragraph with the style 'List Number'
                paragraph = document.add_paragraph(style='List Bullet 2')
                set_paragraph_font(paragraph, font_name="Times New Roman", font_size=11)
                add_hyperlink(paragraph, f"{jira_url}/browse/{row.Issue_key}", f"{row.Issue_key} - {row.Summary}",
                              font_name="Times New Roman", font_size=11)

    # Epic Progress: Resolved tasks by epic
    document.add_heading("Epic Progress", level=2)
    if epic_summary:
        for epic in epic_summary:
            document.add_heading(epic["Epic"], level=3)
            for task in epic["Tasks"]:
                paragraph = document.add_paragraph(f"{task['Task_Key']}: {task['Task_Summary']}", style="List Bullet 2")
                set_paragraph_font(paragraph, font_name="Calibri (Body)", font_size=10)
    else:
        document.add_paragraph("No resolved tasks for open epics during the specified period.")

    # Resolved Tasks: Detailed list by week
    add_resolved_tasks_section(document, resolved_data)

    # Save the Word document
    output_file = f"{file_suffix}.docx"
    document.save(output_file)
    print(f"Word report successfully created: {output_file}")


def generate_report(data, start_date, end_date, project, jira_url, include_empty_weeks, member_list_file=None, pr_stat_file=None):
    """
    Generate Excel and Word reports with team and role-based metrics.

    Args:
        data (pd.DataFrame): Jira data.
        start_date (str): Report start date (YYYY-MM-DD).
        end_date (str): Report end date (YYYY-MM-DD).
        project (str): Jira project key.
        jira_url (str): Jira base URL.
        include_empty_weeks (bool): Include weeks with no activity.
        member_list_file (str, optional): Path to member list Excel file.
        pr_stat_file (str, optional): Path to PR statistics Excel file.
    """
    from team_performance import calculate_role_metrics, export_role_metrics_to_excel, add_role_metrics_to_docx

    start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
    start_monday = start_date - timedelta(days=start_date.weekday())
    valid_weeks = pd.date_range(start=start_monday, end=end_date, freq='W-MON').strftime("%G-W%V").tolist()

    # Filter data to include only valid weeks
    data = data[data["Week"].isin(valid_weeks)]

    if member_list_file:
        required_assignees = read_member_list(member_list_file)
    else:
        required_assignees = data["Assignee"].unique().tolist()

    if include_empty_weeks:
        data = fill_missing_weeks(data, valid_weeks, required_assignees)

    headers = generate_week_headers(valid_weeks, data)
    # Generate epic report data
    epic_summary = generate_epic_report(data)
    # generate team performance
    team_metrics = calculate_team_performance(data, required_assignees)

    start_date = start_date.strftime("%Y-%m-%d")
    file_suffix = generate_file_suffix()
    output_file = f"jira_report_{project}_{start_date}-{end_date}{file_suffix}.docx"
    export_team_performance_to_excel(team_metrics, output_file)
    add_team_performance_to_docx(team_metrics, output_file)

    if pr_stat_file:
        pr_stats_df = pd.read_excel(pr_stat_file)
    else:
        pr_stats_df = pd.DataFrame()

    if member_list_file:
        required_members = pd.read_excel(member_list_file)
    else:
        required_members = pd.DataFrame(columns=["name", "role", "gitee_account"])

    role_metrics_df, team_avg = calculate_role_metrics(data, required_members, pr_stats_df)
    export_role_metrics_to_excel(role_metrics_df, team_avg, output_file)
    add_role_metrics_to_docx(role_metrics_df, team_avg, output_file)

    generate_excel_report(data, start_date, end_date, project, headers, output_file)
    generate_word_report(data, start_date, end_date, project, headers, output_file, jira_url, epic_summary, member_list_file)

def generate_file_suffix():
    """
    Generate a unique file suffix based on the current timestamp.

    Returns:
        str: Timestamp-based suffix (e.g., _20250806174123).
    """
    return datetime.now().strftime("_%Y%m%d%H%M%S")

def main():
    """
    Main function to fetch Jira data and generate reports.

    Requirements:
        - Config file (config.ini) or command-line arguments must provide valid Jira credentials.
        - Member list file (if provided) must have columns: name, role, gitee_account.
        - PR statistics file (if provided) must have columns: Login, Additions, Deletions, Reviewers.
    """
    jira_url, jira_username, jira_password, project, start_date, end_date, include_empty_weeks, member_list_file, pr_stat_file = parse_arguments_and_config()
    jira_options = {"verify": "bundle-ca"} if os.path.exists("bundle-ca") else True
    jira = JIRA(server=jira_url, basic_auth=(jira_username, jira_password), options=jira_options)
    data = fetch_jira_data(jira, project, start_date, end_date)
    generate_report(data, start_date, end_date, project, jira_url, include_empty_weeks, member_list_file, pr_stat_file)

    # TODO:
    # 1. Generate report only for members specified in member_list_file.
    # 2. Support JQL query: worklogAuthor in (members) AND worklogDate >= start_date.

if __name__ == "__main__":
    main()
