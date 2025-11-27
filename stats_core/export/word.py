"""
Word export helpers.
"""

from pathlib import Path
from typing import Iterable, Sequence, Optional

from docx import Document
from docx.table import _Row, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def render_table(document: Document, headers: Sequence[str], rows: Iterable[Sequence[str]], font_name: str = "Calibri (Body)", font_size: int = 8, table_style: str = "Table Grid") -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = table_style
    _apply_row_content(table.rows[0], headers, font_name, font_size, header=True)
    for row in rows:
        row_cells = table.add_row().cells
        _apply_row_content(row_cells, row, font_name, font_size, header=False)


def export_report(output_path: Path, sections: list[dict], template: Optional[Path] = None) -> None:
    document = Document(template) if template and template.exists() else Document()
    for section in sections:
        heading_level = section.get("level", 1)
        title = section.get("title")
        if title:
            document.add_heading(title, level=heading_level)
        render_table(
            document,
            section.get("headers", []),
            section.get("rows", []),
            font_name=section.get("font_name", "Calibri (Body)"),
            font_size=section.get("font_size", 8),
            table_style=section.get("table_style", "Table Grid"),
        )
    document.save(output_path)


def _apply_row_content(row_or_cells, values: Sequence[str], font_name: str, font_size: int, header: bool) -> None:
    cells = row_or_cells.cells if isinstance(row_or_cells, _Row) else row_or_cells
    for idx, value in enumerate(values):
        cell: _Cell = cells[idx]
        cell.text = "" if value is None else str(value)
        _apply_paragraph_style(cell.paragraphs, font_name, font_size, bold=header)


def _apply_paragraph_style(paragraphs: Iterable[Paragraph], font_name: str, font_size: int, bold: bool = False) -> None:
    for paragraph in paragraphs:
        if not paragraph.runs:
            paragraph.add_run("")
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.find(qn("w:rFonts"))
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)
            r_fonts.set(qn("w:ascii"), font_name)
            r_fonts.set(qn("w:hAnsi"), font_name)

