"""
Jira Engineer Weekly Activity report - per engineer, per week activity with comments and time.
"""

from __future__ import annotations

from datetime import timedelta
import re

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..export.word import add_hyperlink
from ..utils.members import read_member_list
from .jira_utils import get_valid_weeks, norm_name


_URL_PATTERN = re.compile(r'https?://[^\s<>"{}|\\^`\[\]]+')


def _format_duration(seconds: int | float | None) -> str:
    if not seconds:
        return "0m"
    total_minutes = int(round(float(seconds) / 60))
    hours = total_minutes // 60
    minutes = total_minutes % 60
    if hours and minutes:
        return f"{hours}h {minutes}m"
    if hours:
        return f"{hours}h"
    return f"{minutes}m"


def _extract_urls(text: str | None) -> list[str]:
    if not text:
        return []
    return _URL_PATTERN.findall(text)


def _clean_comment_body(text: str | None) -> str:
    if not text:
        return ""
    lines = [line.strip() for line in str(text).splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def add_engineer_weekly_activity_to_document(
    document: Document,
    worklogs_df: pd.DataFrame,
    comments_df: pd.DataFrame,
    start_date: str,
    end_date: str,
    jira_url: str,
    member_list_file: str | None = None,
    include_empty: bool = True,
) -> None:
    """
    Add Engineer Weekly Activity section to Word document.

    Args:
        document: Word document to add section to
        worklogs_df: DataFrame with Issue_key, Summary, Assignee, Week, WorklogSeconds
        comments_df: DataFrame with Issue_key, Summary, CommentAuthor, Week, CommentBody
        start_date: Start date string (YYYY-MM-DD)
        end_date: End date string (YYYY-MM-DD)
        jira_url: Base Jira URL for hyperlinks
        member_list_file: Optional path to Excel file with member list
        include_empty: Whether to include weeks with no activity
    """
    document.add_heading("Engineer Weekly Activity", level=1)

    if worklogs_df is None or worklogs_df.empty:
        worklogs_df = pd.DataFrame(columns=[
            "Issue_key",
            "Summary",
            "Assignee",
            "Assignee_norm",
            "Week",
            "WorklogSeconds",
            "Status",
            "Resolution",
        ])
    if comments_df is None or comments_df.empty:
        comments_df = pd.DataFrame(columns=[
            "Issue_key",
            "Summary",
            "CommentAuthor",
            "CommentAuthor_norm",
            "Week",
            "CommentBody",
            "CommentDate",
            "CommentDateStr",
            "Status",
            "Resolution",
        ])

    if "Assignee_norm" not in worklogs_df.columns:
        worklogs_df["Assignee_norm"] = worklogs_df.get("Assignee", "").map(norm_name)
    if "CommentAuthor_norm" not in comments_df.columns:
        comments_df["CommentAuthor_norm"] = comments_df.get("CommentAuthor", "").map(norm_name)

    if member_list_file:
        required_assignees = read_member_list(member_list_file)
    else:
        required_assignees = sorted(
            set(worklogs_df.get("Assignee", pd.Series([], dtype=object)).dropna().tolist())
            | set(comments_df.get("CommentAuthor", pd.Series([], dtype=object)).dropna().tolist())
        )

    valid_weeks = get_valid_weeks(start_date, end_date)

    for assignee in required_assignees:
        assignee_norm = norm_name(assignee)

        paragraph_assignee = document.add_paragraph()
        paragraph_assignee_format = paragraph_assignee.paragraph_format
        paragraph_assignee_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_assignee_format.space_before = Pt(12)
        paragraph_assignee_format.space_after = Pt(12)
        paragraph_assignee_format.line_spacing = 1.0
        paragraph_assignee.style = "Heading 2"

        assignee_run = paragraph_assignee.add_run(assignee)
        assignee_run.font.name = "Times New Roman"
        assignee_run.font.size = Pt(11)
        assignee_run.font.bold = True
        assignee_run.font.color.rgb = RGBColor(0, 0, 0)

        if include_empty:
            weeks_to_show = valid_weeks
        else:
            worklog_weeks = worklogs_df[worklogs_df["Assignee_norm"] == assignee_norm]["Week"].unique().tolist()
            comment_weeks = comments_df[comments_df["CommentAuthor_norm"] == assignee_norm]["Week"].unique().tolist()
            weeks_to_show = sorted(set(worklog_weeks + comment_weeks))

        for week in weeks_to_show:
            year, week_num = map(int, week.split("-W"))
            week_start = pd.Timestamp.fromisocalendar(year, week_num, 1).strftime("%Y-%m-%d")
            week_end = (pd.Timestamp.fromisocalendar(year, week_num, 1) + timedelta(days=6)).strftime("%Y-%m-%d")
            week_header = f"ww{week_num} {week_start}-{week_end}"

            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph_format.space_before = Pt(13)
            paragraph_format.space_after = Pt(13)
            paragraph_format.line_spacing = 1.73
            paragraph.style = "Heading 3"

            paragraph_run = paragraph.add_run(week_header)
            paragraph_run.font.name = "Times New Roman"
            paragraph_run.font.size = Pt(11)
            paragraph_run.font.bold = True
            paragraph_run.font.color.rgb = RGBColor(0, 0, 0)

            worklog_week = worklogs_df[
                (worklogs_df["Assignee_norm"] == assignee_norm) & (worklogs_df["Week"] == week)
            ]
            comment_week = comments_df[
                (comments_df["CommentAuthor_norm"] == assignee_norm) & (comments_df["Week"] == week)
            ]

            issue_keys = sorted(set(worklog_week.get("Issue_key", [])) | set(comment_week.get("Issue_key", [])))

            if not issue_keys:
                paragraph = document.add_paragraph(style="List Bullet 2")
                vacation_run = paragraph.add_run("vacation")
                vacation_run.font.name = "Times New Roman"
                vacation_run.font.size = Pt(11)
                vacation_run.font.color.rgb = RGBColor(0, 0, 0)
                continue

            weekly_links: list[str] = []

            for issue_key in issue_keys:
                issue_summary = ""
                summary_from_worklog = worklog_week[worklog_week["Issue_key"] == issue_key].get("Summary")
                if summary_from_worklog is not None and not summary_from_worklog.empty:
                    issue_summary = str(summary_from_worklog.iloc[0])
                else:
                    summary_from_comment = comment_week[comment_week["Issue_key"] == issue_key].get("Summary")
                    if summary_from_comment is not None and not summary_from_comment.empty:
                        issue_summary = str(summary_from_comment.iloc[0])

                issue_status = ""
                issue_resolution = ""
                status_from_worklog = worklog_week[worklog_week["Issue_key"] == issue_key].get("Status")
                if status_from_worklog is not None and not status_from_worklog.empty:
                    issue_status = str(status_from_worklog.iloc[0])
                else:
                    status_from_comment = comment_week[comment_week["Issue_key"] == issue_key].get("Status")
                    if status_from_comment is not None and not status_from_comment.empty:
                        issue_status = str(status_from_comment.iloc[0])

                resolution_from_worklog = worklog_week[worklog_week["Issue_key"] == issue_key].get("Resolution")
                if resolution_from_worklog is not None and not resolution_from_worklog.empty:
                    issue_resolution = str(resolution_from_worklog.iloc[0])
                else:
                    resolution_from_comment = comment_week[comment_week["Issue_key"] == issue_key].get("Resolution")
                    if resolution_from_comment is not None and not resolution_from_comment.empty:
                        issue_resolution = str(resolution_from_comment.iloc[0])

                total_seconds = worklog_week[worklog_week["Issue_key"] == issue_key]["WorklogSeconds"].sum()
                time_str = _format_duration(total_seconds)

                issue_paragraph = document.add_paragraph(style="List Bullet 2")
                if isinstance(issue_key, str) and issue_key.strip():
                    add_hyperlink(
                        issue_paragraph,
                        f"{jira_url}/browse/{issue_key}",
                        f"{issue_key} - {issue_summary}",
                        font_name="Times New Roman",
                        font_size=11
                    )
                else:
                    issue_paragraph.add_run(f"{issue_key} - {issue_summary}")
                status_parts = []
                if issue_status:
                    status_parts.append(f"Status: {issue_status}")
                if issue_resolution:
                    status_parts.append(f"Resolution: {issue_resolution}")
                if status_parts:
                    issue_paragraph.add_run(f" ({', '.join(status_parts)})")
                issue_paragraph.add_run(f" (time: {time_str})")

                issue_comments = comment_week[comment_week["Issue_key"] == issue_key].copy()
                if not issue_comments.empty:
                    issue_comments = issue_comments.sort_values(by="CommentDate")
                    for _, comment_row in issue_comments.iterrows():
                        comment_body = _clean_comment_body(comment_row.get("CommentBody", ""))
                        if not comment_body:
                            continue
                        comment_date = comment_row.get("CommentDateStr") or ""
                        comment_author = comment_row.get("CommentAuthor") or ""
                        comment_paragraph = document.add_paragraph(style="List Bullet 3")
                        header_run = comment_paragraph.add_run(f"[{comment_date}] {comment_author}:")
                        header_run.add_break()
                        comment_paragraph.add_run(comment_body)
                        weekly_links.extend(_extract_urls(comment_body))

            if weekly_links:
                links_header = document.add_paragraph(style="List Bullet 2")
                links_header.add_run("Links:")
                for link in weekly_links:
                    link_paragraph = document.add_paragraph(style="List Bullet 3")
                    link_paragraph.add_run(link)
