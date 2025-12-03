"""
Jira Table View report - tabular format with Name, Week #, Date, Description, Link, Status.
"""

from __future__ import annotations

from datetime import timedelta

import pandas as pd
from docx import Document

from ..export.word import add_hyperlink
from ..utils.members import read_member_list
from .jira_utils import norm_name, is_empty_task


def add_table_view_to_document(
    document: Document,
    data: pd.DataFrame,
    jira_url: str,
    member_list_file: str | None = None,
) -> None:
    """
    Add Table View section to Word document - tabular format.

    Args:
        document: Word document to add section to
        data: DataFrame with Issue_key, Summary, Assignee, Status, Week columns
        jira_url: Base Jira URL for hyperlinks
        member_list_file: Optional path to Excel file with member list
    """
    document.add_heading("Tabular View", level=2)
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Add headers
    headers = ["Name", "Week #", "Date", "Description", "Link", "Status"]
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header

    # Get required assignees
    if member_list_file:
        required_assignees = read_member_list(member_list_file)
    else:
        required_assignees = sorted(data["Assignee"].unique())

    # Sort data by assignee, week and status
    sorted_data = data.sort_values(by=["Assignee", "Week", "Status"])

    # Add data rows, ensuring all required assignees are included
    for assignee in required_assignees:
        assignee_data = sorted_data[sorted_data["Assignee_norm"] == norm_name(assignee)]

        if assignee_data.empty:
            # Add a row with empty task details if no data for the assignee
            row_cells = table.add_row().cells
            row_cells[0].text = assignee
            row_cells[1].text = ""
            row_cells[2].text = ""
            row_cells[3].text = "vacation"
            row_cells[4].text = ""
            row_cells[5].text = ""
        else:
            # Add rows for each task of the assignee
            for _, row in assignee_data.iterrows():
                year, week_num = map(int, row["Week"].split("-W"))
                week_start = pd.Timestamp.fromisocalendar(year, week_num, 1).strftime("%Y/%m/%d")
                week_end = (pd.Timestamp.fromisocalendar(year, week_num, 1) + timedelta(days=6)).strftime("%Y/%m/%d")
                week_range = f"{week_start} – {week_end}"

                # Add a row to the table
                row_cells = table.add_row().cells
                row_cells[0].text = assignee
                row_cells[1].text = row["Week"]
                row_cells[2].text = week_range
                
                # Clear cell before adding hyperlink
                cell_paragraph = row_cells[4].paragraphs[0]
                cell_paragraph.clear()
                
                if is_empty_task(row["Summary"], row["Status"]):
                    row_cells[3].text = "vacation"
                    row_cells[5].text = ""
                else:
                    desc = row["Summary"]
                    if row.get("Reassigned", False):
                        desc = "[reassigned] " + desc
                    row_cells[3].text = desc
                    row_cells[5].text = row["Status"]
                    
                    if isinstance(row["Issue_key"], str) and row["Issue_key"].strip():
                        add_hyperlink(
                            cell_paragraph,
                            f"{jira_url}/browse/{row['Issue_key']}",
                            f"{row['Issue_key']}",
                            font_size=8
                        )

    # Apply font formatting to all cells
    from ..export.word import _apply_paragraph_style
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _apply_paragraph_style([paragraph], font_name="Calibri (Body)", font_size=8)

