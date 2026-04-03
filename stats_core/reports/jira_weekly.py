"""
Jira weekly report - coordinates List View, Table View, Epic Progress, and Excel export.
"""

from __future__ import annotations

import json
from datetime import datetime
import logging
from pathlib import Path
from configparser import ConfigParser
import re
from typing import Any

import pandas as pd
from docx import Document

from . import registry
from .jira_comprehensive import rewrite_summary_items_with_ai
from .jira_utils import (
    _comment_body_to_text,
    fetch_jira_data,
    fetch_jira_activity_data,
    build_resolved_issues_snapshot,
    mark_reassigned_tasks,
    fill_missing_weeks,
    generate_week_headers,
    get_valid_weeks,
)
from .jira_list_view import add_list_view_to_document
from .jira_engineer_weekly import add_engineer_weekly_activity_to_document
from .jira_table_view import add_table_view_to_document
from .jira_epic_report import (
    generate_epic_resolved_hierarchy,
    generate_epic_progress_from_worklogs,
    add_epic_progress_to_document,
    add_resolved_tasks_section,
)
from ..sources.jira import JiraSource
from ..utils.members import read_member_list
from ..utils.parallel import parallel_map
from ..utils.progress import NoopProgressManager

logger = logging.getLogger(__name__)

_SUMMARY_COLUMNS = [
    "Epic_Link",
    "Epic_Name",
    "Summary",
    "Planned_Tasks_Resolved",
    "Reported_Issues_Resolved",
]
_SUBTASK_TYPES = {"sub-task", "subtask", "sub task"}
_FILELIKE_PATTERN = re.compile(
    r"\b[\w.\- ]+\.(?:png|jpe?g|gif|bmp|webp|svg|pdf|docx?|xlsx?|pptx?|txt|log|json|ya?ml|xml|csv|zip|rar|7z)\b",
    re.IGNORECASE,
)
_JIRA_BLOCK_MACRO_PATTERN = re.compile(
    r"\{(?:code(?::[^}]*)?|noformat)\}.*?\{(?:code|noformat)\}",
    re.IGNORECASE | re.DOTALL,
)


def _compact_text(value: Any) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    return " ".join(str(value).strip().split())


def _normalize_text(value: Any) -> str:
    return _compact_text(value).casefold()


def _to_plain_text(value: Any) -> str:
    if value is None:
        return ""
    return _compact_text(_comment_body_to_text(value))


def _is_unknown_epic_name(value: Any) -> bool:
    return _normalize_text(value) in {"", "unknown epic"}


def _is_countable_resolution(value: Any) -> bool:
    resolution_norm = _normalize_text(value)
    return not any(marker in resolution_norm for marker in ("won't do", "wont do", "invalid"))


def _sanitize_weekly_summary_evidence(value: Any) -> str:
    cleaned = _to_plain_text(value)
    if not cleaned:
        return ""
    cleaned = _JIRA_BLOCK_MACRO_PATTERN.sub(" ", cleaned)
    cleaned = re.sub(r"\{(?:code(?::[^}]*)?|noformat|quote|panel|color(?::[^}]*)?)\}", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"!\[[^\]]*\]\([^)]+\)", " ", cleaned)
    cleaned = re.sub(r"!\[[^\]]*\]", " ", cleaned)
    cleaned = re.sub(r"\[[^\]]+\]\([^)]+\)", " ", cleaned)
    cleaned = re.sub(r"(?:https?://|ftp://|file://|www\.)\S+", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\\\\[^\s]+", " ", cleaned)
    cleaned = re.sub(r"\b[A-Za-z]:\\[^\s]+", " ", cleaned)
    cleaned = re.sub(r"(?<![A-Za-z0-9])/(?:[\w.\-]+/)+[\w.\-]+", " ", cleaned)
    cleaned = re.sub(r"\[\^[^\]\r\n]+\]", " ", cleaned)
    cleaned = re.sub(r"<[^>\r\n]+>", " ", cleaned)
    cleaned = _FILELIKE_PATTERN.sub(" ", cleaned)
    cleaned = re.sub(r"\b[A-Z]+-\d+\b", " ", cleaned)
    cleaned = re.sub(r"\b[0-9a-f]{7,40}\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"(?i)\b(?:pr|mr|pull request|merge request|commit)\b\s*[:#-]?\s*[A-Za-z0-9/_-]*", " ", cleaned)
    cleaned = re.sub(r"[`*_>#]+", " ", cleaned)
    cleaned = re.sub(r"(?<!\w)[\[\]{}()]+(?=\s|$)", " ", cleaned)
    cleaned = re.sub(r"(?<!\w)(?:!\[|\[!)+", " ", cleaned)
    cleaned = " ".join(
        token
        for token in cleaned.split()
        if not re.fullmatch(r"[\[\]{}()!|\\/:;.,+-]+", token)
    )
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" -,:;")
    return cleaned


def _limit_summary_text(text: Any, *, max_sentences: int = 2, max_words: int = 45) -> str:
    cleaned = _sanitize_weekly_summary_evidence(text)
    if not cleaned:
        return ""
    sentences = [
        chunk.strip(" -,:;")
        for chunk in re.split(r"(?<=[.!?])\s+", cleaned)
        if chunk.strip(" -,:;")
    ]
    if not sentences:
        return ""
    limited = " ".join(sentences[:max_sentences]).strip()
    words = limited.split()
    if len(words) > max_words:
        limited = " ".join(words[:max_words]).rstrip(" ,;:-")
        if limited and limited[-1] not in ".!?":
            limited += "."
    elif limited and limited[-1] not in ".!?":
        limited += "."
    return limited


def _append_unique_text(target: list[str], values: list[str]) -> None:
    seen = {_normalize_text(value) for value in target}
    for value in values:
        cleaned = _compact_text(value)
        if not cleaned:
            continue
        marker = _normalize_text(cleaned)
        if marker in seen:
            continue
        seen.add(marker)
        target.append(cleaned)


def _extract_comment_fact_map(comments_df: pd.DataFrame) -> dict[str, list[str]]:
    if comments_df is None or comments_df.empty:
        return {}

    filtered_df = comments_df.copy()
    if "Is_Worklog_Comment" in filtered_df.columns:
        filtered_df = filtered_df[~filtered_df["Is_Worklog_Comment"].fillna(False).astype(bool)]
    if filtered_df.empty:
        return {}

    sort_columns = [column for column in ("Issue_key", "CommentDate", "CommentId") if column in filtered_df.columns]
    if sort_columns:
        filtered_df = filtered_df.sort_values(by=sort_columns, kind="mergesort")

    fact_map: dict[str, list[str]] = {}
    for issue_key, issue_comments in filtered_df.groupby("Issue_key", dropna=False, sort=True):
        normalized_key = _compact_text(issue_key)
        if not normalized_key:
            continue
        facts: list[str] = []
        for _, row in issue_comments.iterrows():
            cleaned = _sanitize_weekly_summary_evidence(row.get("CommentBody"))
            if not cleaned:
                continue
            fragments = [
                _compact_text(fragment).strip(" -,:;")
                for fragment in re.split(r"(?:\n+|(?<=[.!?])\s+)", cleaned)
                if _compact_text(fragment).strip(" -,:;")
            ]
            normalized_facts: list[str] = []
            for fragment in fragments:
                if len(fragment) < 8:
                    continue
                fragment_norm = _normalize_text(fragment)
                if fragment_norm.startswith("see ") and len(fragment.split()) <= 4:
                    continue
                if fragment[-1] not in ".!?":
                    fragment += "."
                normalized_facts.append(fragment)
                if len(normalized_facts) >= 4:
                    break
            _append_unique_text(facts, normalized_facts)
            if len(facts) >= 6:
                break
        fact_map[normalized_key] = facts
    return fact_map


def _fetch_parent_details(
    jira_source: JiraSource,
    resolved_df: pd.DataFrame,
) -> dict[str, dict[str, str]]:
    issue_keys = {
        _compact_text(issue_key)
        for issue_key in resolved_df.get("Issue_key", pd.Series(dtype=object)).fillna("").astype(str).tolist()
        if _compact_text(issue_key)
    }
    parent_keys: set[str] = set()
    for _, row in resolved_df.iterrows():
        row_type = _normalize_text(row.get("Type"))
        parent_key = _compact_text(row.get("Parent")) or _compact_text(row.get("Parent_Key"))
        if row_type in _SUBTASK_TYPES and parent_key and parent_key not in issue_keys:
            parent_keys.add(parent_key)
    if not parent_keys:
        return {}
    try:
        details = jira_source.fetch_issue_details(sorted(parent_keys))
    except Exception as exc:
        logger.warning("Weekly summary: failed to fetch parent details: %s", exc)
        return {}
    return details if isinstance(details, dict) else {}


def _build_epic_name_map(
    jira_source: JiraSource,
    resolved_df: pd.DataFrame,
    parent_details: dict[str, dict[str, str]],
) -> dict[str, str]:
    epic_names: dict[str, str] = {}
    missing_epics: set[str] = set()

    for _, row in resolved_df.iterrows():
        epic_link = _compact_text(row.get("Epic_Link"))
        epic_name = _compact_text(row.get("Epic_Name"))
        if not epic_link:
            continue
        if epic_name and not _is_unknown_epic_name(epic_name):
            epic_names[epic_link] = epic_name
        else:
            missing_epics.add(epic_link)

    for detail in parent_details.values():
        epic_link = _compact_text(detail.get("Epic_Link"))
        if epic_link and epic_link not in epic_names:
            missing_epics.add(epic_link)

    missing_epics -= set(epic_names)
    if missing_epics:
        try:
            epic_names.update(jira_source.fetch_epic_names(sorted(missing_epics)))
        except Exception as exc:
            logger.warning("Weekly summary: failed to fetch epic names: %s", exc)
    return epic_names


def _build_weekly_summary_fallback(
    anchor_title: str,
    anchor_resolved: bool,
    resolved_items: list[str],
    comment_facts: list[str],
    anchor_description: str,
) -> str:
    base = _sanitize_weekly_summary_evidence(anchor_title) or "Task group"
    delivered_items = [
        _sanitize_weekly_summary_evidence(item)
        for item in resolved_items
        if _sanitize_weekly_summary_evidence(item)
        and _normalize_text(_sanitize_weekly_summary_evidence(item)) != _normalize_text(base)
    ]
    if delivered_items:
        if len(delivered_items) == 1:
            first = f"{base}: completed {delivered_items[0]}."
        elif len(delivered_items) == 2:
            first = f"{base}: completed {delivered_items[0]} and {delivered_items[1]}."
        else:
            first = f"{base}: completed {delivered_items[0]}, {delivered_items[1]}, and additional scoped work."
    elif anchor_resolved:
        first = f"{base} was completed."
    else:
        first = f"{base}: delivered progress through resolved subtasks."

    detail_source = comment_facts[0] if comment_facts else anchor_description
    detail = _limit_summary_text(detail_source, max_sentences=1, max_words=18)
    if detail and _normalize_text(detail) != _normalize_text(first):
        return _limit_summary_text(f"{first} {detail}", max_sentences=2, max_words=50)
    return _limit_summary_text(first, max_sentences=2, max_words=50)


def _build_weekly_summary_prompt(
    items: list[dict[str, Any]],
    *,
    start_index: int = 1,
) -> tuple[dict[str, str], str]:
    prompt_lines = [
        "You are preparing a weekly software delivery summary for management.",
        "Task: rewrite EACH task group into concise achievement text.",
        "Input structure:",
        "- epic_name: epic grouping context",
        "- anchor_title: parent task, feature, or bug title",
        "- anchor_type: parent issue type",
        "- anchor_status: current parent issue status",
        "- anchor_description: sanitized parent task context",
        "- resolved_items: completed tasks and subtasks for this group in the selected period",
        "- comment_facts: sanitized factual statements extracted from period comments across the group",
        "Strict rules:",
        "1) Output language: English.",
        "2) Treat resolved subtasks as achievements of the parent task.",
        "3) Describe what was delivered during the selected period for the whole task group.",
        "4) Use only provided facts; do not invent details.",
        "5) Ignore links, Jira keys, PR/MR mentions, repository references, commit hashes, file names, uploaded artifact names, absolute paths, and UNC paths.",
        "6) Do not mention where evidence was stored.",
        "7) Write 1-2 short complete sentences, factual and report-ready.",
        "8) Return ONLY one valid JSON object mapping id to rewritten text.",
        "JSON example: {\"t1\":\"...\", \"t2\":\"...\"}",
        "---",
        "Input items:",
    ]
    target_map: dict[str, str] = {}
    for idx, item in enumerate(items, start=start_index):
        target_id = f"t{idx}"
        target_map[target_id] = str(item["id"])
        prompt_lines.append(
            "ID={target_id}; epic_name={epic_name}; anchor_title={anchor_title}; "
            "anchor_type={anchor_type}; anchor_status={anchor_status}; anchor_description={anchor_description}; "
            "resolved_items={resolved_items}; comment_facts={comment_facts}".format(
                target_id=target_id,
                epic_name=json.dumps(str(item.get("epic_name", ""))),
                anchor_title=json.dumps(str(item.get("anchor_title", ""))),
                anchor_type=json.dumps(str(item.get("anchor_type", ""))),
                anchor_status=json.dumps(str(item.get("anchor_status", ""))),
                anchor_description=json.dumps(str(item.get("anchor_description", ""))),
                resolved_items=json.dumps(" | ".join(str(value) for value in item.get("resolved_items", []))),
                comment_facts=json.dumps(" | ".join(str(value) for value in item.get("comment_facts", []))),
            )
        )
    return target_map, "\n".join(prompt_lines)


def _sanitize_weekly_summary_ai_text(text: Any) -> str:
    return _limit_summary_text(text, max_sentences=2, max_words=50)


def build_weekly_epic_summary_df(
    jira_source: JiraSource,
    resolved_issues_df: pd.DataFrame,
    comments_df: pd.DataFrame,
    start_date: str,
    end_date: str,
    config: ConfigParser,
    extra_params: dict[str, Any],
) -> pd.DataFrame:
    if resolved_issues_df is None or resolved_issues_df.empty:
        return pd.DataFrame(columns=_SUMMARY_COLUMNS)

    resolved_df = resolved_issues_df.copy()
    for column in (
        "Issue_key",
        "Summary",
        "Type",
        "Epic_Link",
        "Epic_Name",
        "Status",
        "Resolution",
        "Resolution_Date",
        "Parent",
        "Parent_Key",
        "Parent_Summary",
        "Description",
    ):
        if column not in resolved_df.columns:
            resolved_df[column] = ""
    resolved_df["Parent"] = resolved_df["Parent"].fillna("").astype(str)
    resolved_df["Parent_Key"] = resolved_df["Parent_Key"].fillna("").astype(str)
    empty_parent_mask = resolved_df["Parent"].str.strip().eq("")
    resolved_df.loc[empty_parent_mask, "Parent"] = resolved_df.loc[empty_parent_mask, "Parent_Key"]
    resolved_df["_epic_sort"] = resolved_df["Epic_Name"].fillna("").astype(str).str.casefold()
    resolved_df["_resolved_sort"] = pd.to_datetime(resolved_df["Resolution_Date"], errors="coerce")
    resolved_df = resolved_df.sort_values(
        by=["_epic_sort", "_resolved_sort", "Issue_key"],
        ascending=[True, True, True],
        na_position="last",
        kind="mergesort",
    )

    parent_details = _fetch_parent_details(jira_source, resolved_df)
    epic_name_map = _build_epic_name_map(jira_source, resolved_df, parent_details)
    comment_fact_map = _extract_comment_fact_map(comments_df)
    issue_row_map = {
        _compact_text(row.get("Issue_key")): row.to_dict()
        for _, row in resolved_df.iterrows()
        if _compact_text(row.get("Issue_key"))
    }

    groups: dict[tuple[str, str, str], dict[str, Any]] = {}
    group_order: list[tuple[str, str, str]] = []
    epic_order: list[tuple[str, str]] = []
    epic_counters: dict[tuple[str, str], dict[str, int]] = {}

    for _, row in resolved_df.iterrows():
        issue_key = _compact_text(row.get("Issue_key"))
        if not issue_key or not _is_countable_resolution(row.get("Resolution")):
            continue

        issue_type = _compact_text(row.get("Type")) or "Task"
        type_norm = _normalize_text(issue_type)
        if type_norm == "epic":
            continue

        parent_key = _compact_text(row.get("Parent")) or _compact_text(row.get("Parent_Key"))
        is_subtask = type_norm in _SUBTASK_TYPES
        anchor_key = parent_key if is_subtask and parent_key else issue_key

        anchor_row = issue_row_map.get(anchor_key, {})
        parent_detail = parent_details.get(anchor_key, {})
        anchor_title = _sanitize_weekly_summary_evidence(
            parent_detail.get("Summary")
            or anchor_row.get("Summary")
            or row.get("Parent_Summary")
            or row.get("Summary")
        )
        anchor_type = _compact_text(parent_detail.get("Type") or anchor_row.get("Type") or issue_type or "Task")
        anchor_status = _compact_text(parent_detail.get("Status") or anchor_row.get("Status") or row.get("Status"))
        anchor_description = _sanitize_weekly_summary_evidence(
            parent_detail.get("Description") or anchor_row.get("Description") or ""
        )
        anchor_epic_link = _compact_text(
            row.get("Epic_Link")
            or anchor_row.get("Epic_Link")
            or parent_detail.get("Epic_Link")
        )
        anchor_epic_name = _compact_text(row.get("Epic_Name") or anchor_row.get("Epic_Name"))
        if (not anchor_epic_name or _is_unknown_epic_name(anchor_epic_name)) and anchor_epic_link:
            anchor_epic_name = _compact_text(epic_name_map.get(anchor_epic_link))
        if not anchor_epic_name:
            anchor_epic_name = "Unknown Epic"

        epic_key = (anchor_epic_link, anchor_epic_name)
        if epic_key not in epic_counters:
            epic_counters[epic_key] = {"planned_count": 0, "bug_count": 0}
            epic_order.append(epic_key)
        if type_norm == "bug":
            epic_counters[epic_key]["bug_count"] += 1
        else:
            epic_counters[epic_key]["planned_count"] += 1

        group_key = (anchor_epic_link, anchor_epic_name, anchor_key)
        if group_key not in groups:
            groups[group_key] = {
                "id": f"{anchor_epic_link or anchor_epic_name}::{anchor_key}",
                "epic_link": anchor_epic_link,
                "epic_name": anchor_epic_name,
                "anchor_key": anchor_key,
                "anchor_title": anchor_title or _sanitize_weekly_summary_evidence(row.get("Summary")),
                "anchor_type": anchor_type,
                "anchor_status": anchor_status,
                "anchor_description": anchor_description,
                "anchor_resolved": anchor_key == issue_key,
                "resolved_items": [],
                "resolved_item_titles": [],
                "comment_issue_keys": [],
            }
            group_order.append(group_key)

        group = groups[group_key]
        if anchor_key == issue_key:
            group["anchor_resolved"] = True
            if anchor_title:
                group["anchor_title"] = anchor_title
            if anchor_type:
                group["anchor_type"] = anchor_type
            if anchor_status:
                group["anchor_status"] = anchor_status
            if anchor_description:
                group["anchor_description"] = anchor_description
        item_title = _sanitize_weekly_summary_evidence(row.get("Summary"))
        label = f"{issue_type}: {item_title}" if issue_type and item_title else item_title
        if label:
            _append_unique_text(group["resolved_items"], [label])
        if item_title:
            _append_unique_text(group["resolved_item_titles"], [item_title])
        _append_unique_text(group["comment_issue_keys"], [anchor_key, issue_key])

    ai_items: list[dict[str, Any]] = []
    for group_key in group_order:
        group = groups[group_key]
        comment_facts: list[str] = []
        for issue_key in group["comment_issue_keys"]:
            _append_unique_text(comment_facts, comment_fact_map.get(issue_key, []))
        group["comment_facts"] = comment_facts
        group["fallback"] = _build_weekly_summary_fallback(
            group["anchor_title"],
            bool(group["anchor_resolved"]),
            group["resolved_item_titles"],
            comment_facts,
            group["anchor_description"],
        )
        ai_items.append(
            {
                "id": group["id"],
                "epic_name": group["epic_name"],
                "anchor_title": group["anchor_title"],
                "anchor_type": group["anchor_type"],
                "anchor_status": group["anchor_status"],
                "anchor_description": group["anchor_description"],
                "resolved_items": group["resolved_items"],
                "comment_facts": comment_facts,
            }
        )

    rewrite_map = rewrite_summary_items_with_ai(
        ai_items,
        config,
        extra_params,
        prompt_builder=_build_weekly_summary_prompt,
        result_sanitizer=_sanitize_weekly_summary_ai_text,
        system_prompt=(
            "You rewrite structured Jira evidence into concise weekly achievement statements. "
            "Return strict JSON only."
        ),
    ) if ai_items else {}

    rows: list[dict[str, Any]] = []
    for epic_key in epic_order:
        epic_link, epic_name = epic_key
        lines: list[str] = []
        for group_key in group_order:
            group = groups[group_key]
            if (group["epic_link"], group["epic_name"]) != epic_key:
                continue
            rewritten = _sanitize_weekly_summary_ai_text(rewrite_map.get(group["id"]))
            if not rewritten:
                rewritten = _sanitize_weekly_summary_ai_text(group["fallback"]) or group["fallback"]
            if rewritten:
                lines.append(f"- {rewritten}")

        counters = epic_counters.get(epic_key, {"planned_count": 0, "bug_count": 0})
        lines.append(f"Resolved {int(counters['planned_count'])} planned tasks on time.")
        if int(counters["bug_count"]) > 0:
            lines.append(f"Resolved {int(counters['bug_count'])} reported issues.")
        rows.append(
            {
                "Epic_Link": epic_link,
                "Epic_Name": epic_name,
                "Summary": "\n".join(lines),
                "Planned_Tasks_Resolved": int(counters["planned_count"]),
                "Reported_Issues_Resolved": int(counters["bug_count"]),
            }
        )

    return pd.DataFrame(rows, columns=_SUMMARY_COLUMNS)

def _parse_bool(value: str | bool | None, default: bool) -> bool:
    """Parse boolean value from string or return default."""
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    return value.lower() in {"1", "true", "yes", "y", "on"}


def generate_file_suffix() -> str:
    """Generate a timestamp-based suffix for file names to ensure uniqueness."""
    now = datetime.now()
    return now.strftime("_%Y%m%d_%H%M")


def _get_progress(extra_params: dict[str, Any], total_steps: int):
    progress = extra_params.get("progress_manager")
    if progress is None:
        progress = NoopProgressManager()
    progress.set_total(total_steps)
    return progress


def _parallel_workers(extra_params: dict[str, Any], default: int = 4) -> int:
    raw = extra_params.get("parallel_workers")
    try:
        value = int(str(raw))
        return max(value, 1)
    except Exception:
        return default


def generate_excel_report(
    data: pd.DataFrame,
    start_date: str,
    end_date: str,
    project: str,
    headers: list[str],
    output_file: Path,
) -> None:
    """
    Generate an Excel report summarizing the data grouped by assignee and week.

    Args:
        data: DataFrame with Issue_key, Summary, Assignee, Status, Week columns
        start_date: Start date string (YYYY-MM-DD)
        end_date: End date string (YYYY-MM-DD)
        project: Jira project key
        headers: List of week header strings
        output_file: Output file path (without extension)
    """
    data = data.copy()
    data["Formatted"] = data["Status"] + ": " + data["Issue_key"] + " - " + data["Summary"]

    grouped_data = (
        data.groupby(["Assignee", "Week"])["Formatted"]
        .apply("\n".join)
        .unstack(fill_value="")
    )

    grouped_data.columns = headers

    excel_path = Path(f"{output_file}.xlsx")
    with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
        grouped_data.to_excel(writer, sheet_name="Weekly_Grid")
        weekly_sheet = writer.book["Weekly_Grid"]
        weekly_sheet.freeze_panes = "B2"

    logger.info(
        "Excel report successfully created: %s (sheets: %s)",
        excel_path,
        ["Weekly_Grid"],
    )


def _to_weekly_summary_source_df(resolved_issues_df: pd.DataFrame) -> pd.DataFrame:
    if resolved_issues_df.empty:
        return pd.DataFrame(
            columns=[
                "Issue_Key",
                "Summary",
                "Type",
                "Epic_Link",
                "Epic_Name",
                "Epic_Status",
                "Epic_Resolved",
                "Epic_Labels",
                "Resolved",
                "Status",
                "Resolution",
                "Labels",
                "Parent",
                "Parent_Key",
                "Parent_Summary",
                "Description",
                "Last_Comment",
            ]
        )
    summary_df = resolved_issues_df.rename(
        columns={
            "Issue_key": "Issue_Key",
            "Resolution_Date": "Resolved",
        }
    ).copy()
    for column in (
        "Issue_Key",
        "Summary",
        "Type",
        "Epic_Link",
        "Epic_Name",
        "Epic_Status",
        "Epic_Resolved",
        "Epic_Labels",
        "Resolved",
        "Status",
        "Resolution",
        "Labels",
        "Parent",
        "Parent_Key",
        "Parent_Summary",
    ):
        if column not in summary_df.columns:
            summary_df[column] = ""
    if "Parent_Key" in summary_df.columns:
        summary_df["Parent"] = summary_df["Parent"].fillna("").astype(str)
        parent_keys = summary_df["Parent_Key"].fillna("").astype(str)
        empty_parent_mask = summary_df["Parent"].str.strip().eq("")
        summary_df.loc[empty_parent_mask, "Parent"] = parent_keys.loc[empty_parent_mask]
    if "Description" not in summary_df.columns:
        summary_df["Description"] = ""
    if "Last_Comment" not in summary_df.columns:
        summary_df["Last_Comment"] = ""
    return summary_df


def add_summary_section_to_document(document: Document, summary_df: pd.DataFrame) -> None:
    document.add_heading("Summary", level=1)
    if summary_df.empty:
        document.add_paragraph("No summary items for the specified period.")
        return

    sorted_summary = summary_df.copy()
    sorted_summary["_epic_sort"] = sorted_summary["Epic_Name"].fillna("").astype(str).str.casefold()
    sorted_summary = sorted_summary.sort_values(by=["_epic_sort", "Epic_Link"], kind="mergesort")

    for _, row in sorted_summary.iterrows():
        epic_name = str(row.get("Epic_Name") or "Unknown Epic")
        epic_link = str(row.get("Epic_Link") or "").strip()
        heading = f"{epic_name} ({epic_link})" if epic_link else epic_name
        document.add_heading(heading, level=2)

        summary_text = str(row.get("Summary") or "").strip()
        if not summary_text:
            document.add_paragraph("No summary details.")
            continue
        for line in summary_text.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet 2")
            else:
                document.add_paragraph(line)


@registry.register
class JiraWeeklyReport:
    name = "jira_weekly"

    def run(
        self,
        dataset: dict,
        config: ConfigParser,
        output_formats: list[str],
        extra_params: dict | None = None,
    ) -> None:
        extra_params = extra_params or {}
        output_excel = "excel" in output_formats
        output_word = "word" in output_formats
        total_steps = 2 + (1 if output_excel else 0) + (1 if output_word else 0)
        progress = _get_progress(extra_params, total_steps)

        project = extra_params.get("project") or config.get("jira", "project", fallback=None)
        if not project:
            raise ValueError("Project key is required for Jira weekly report. Pass --params project=ABC.")

        start_date = extra_params.get("start") or extra_params.get("start_date")
        end_date = extra_params.get("end") or extra_params.get("end_date")
        if not start_date or not end_date:
            raise ValueError("start_date and end_date must be provided (use --start/--end).")

        include_empty = _parse_bool(extra_params.get("include_empty_weeks"), True)
        member_list_file = extra_params.get("member_list_file") or extra_params.get("members_file")

        # Initialize Jira source
        jira_source = JiraSource(config["jira"])
        jira_url = jira_source.jira_url

        with progress.step("Fetch Jira data"):
            max_workers = _parallel_workers(extra_params)
            if max_workers > 1:
                def _fetch_main():
                    return fetch_jira_data(jira_source, project, start_date, end_date)

                def _fetch_activity():
                    return fetch_jira_activity_data(jira_source, project, start_date, end_date)

                def _fetch_resolved():
                    return build_resolved_issues_snapshot(jira_source, project, start_date, end_date)

                main_df, activity_pair, resolved_issues_df = parallel_map(
                    lambda fn: fn(),
                    [_fetch_main, _fetch_activity, _fetch_resolved],
                    max_workers=min(max_workers, 3),
                    progress_manager=progress,
                    child_label="Jira fetch",
                )
                data = main_df
                worklogs_df, comments_df = activity_pair
            else:
                data = fetch_jira_data(jira_source, project, start_date, end_date)
                worklogs_df, comments_df = fetch_jira_activity_data(jira_source, project, start_date, end_date)
                resolved_issues_df = build_resolved_issues_snapshot(jira_source, project, start_date, end_date)

        # If there is no JIRA data at all, create an empty frame with expected columns
        # so that downstream utilities (fill_missing_weeks, epic reports) work safely.
        if data.empty:
            data = pd.DataFrame(
                columns=[
                    "Issue_key",
                    "Summary",
                    "Assignee",
                    "Final_Assignee",
                    "Status",
                    "Resolution_Date",
                    "Created_Date",
                    "Week",
                    "Epic_Link",
                    "Epic_Name",
                    "Parent_Key",
                    "Parent_Summary",
                    "Type",
                ]
            )

        with progress.step("Process data"):
            start_dt = datetime.strptime(start_date, "%Y-%m-%d").date()
            end_dt = datetime.strptime(end_date, "%Y-%m-%d").date()
            valid_weeks = get_valid_weeks(start_date, end_date)

        # Update the data to include only valid weeks
        if not data.empty:
            data = data[data["Week"].isin(valid_weeks)]

        if member_list_file:
            required_assignees = read_member_list(member_list_file)
        else:
            required_assignees = data["Assignee"].unique().tolist() if not data.empty else []

        if include_empty:
            data = fill_missing_weeks(data, valid_weeks, required_assignees)
            data = mark_reassigned_tasks(data)

        headers = generate_week_headers(valid_weeks, data)
        epic_summary = generate_epic_resolved_hierarchy(resolved_issues_df)
        epic_progress_summary = generate_epic_progress_from_worklogs(worklogs_df)
        weekly_summary_df = (
            build_weekly_epic_summary_df(
                jira_source,
                resolved_issues_df,
                comments_df,
                start_date,
                end_date,
                config,
                extra_params,
            )
            if output_word
            else pd.DataFrame(columns=_SUMMARY_COLUMNS)
        )

        # Generate file suffix
        file_suffix = generate_file_suffix()
        
        # Все отчеты сохраняются в папку reports по умолчанию
        output_dir = extra_params.get("output_dir") or config.get("reporting", "output_dir", fallback="reports")
        output_base = Path(output_dir)
        output_base.mkdir(parents=True, exist_ok=True)
        
        output_file = output_base / f"jira_report_{project}_{start_date}-{end_date}{file_suffix}"

        # Generate Excel report if requested
        if output_excel:
            with progress.step("Export Excel"):
                generate_excel_report(
                    data,
                    start_date,
                    end_date,
                    project,
                    headers,
                    output_file,
                )

        # Generate Word report if requested
        if output_word:
            with progress.step("Export Word"):
                document = Document()
                document.add_heading(f"JIRA Report: {project} - {start_date}-{end_date}", level=1)

                # Add Table View
                add_table_view_to_document(document, data, jira_url, member_list_file)

                # Add List View
                add_list_view_to_document(document, data, start_date, end_date, jira_url, member_list_file)

                # Add Engineer Weekly Activity
                add_engineer_weekly_activity_to_document(
                    document,
                    worklogs_df,
                    comments_df,
                    start_date,
                    end_date,
                    jira_url,
                    member_list_file,
                    include_empty=include_empty,
                )

                # Add Epic Progress
                add_epic_progress_to_document(document, epic_summary, jira_url, epic_progress_summary)

                # Add Summary section
                add_summary_section_to_document(document, weekly_summary_df)

                # Add Resolved Tasks section
                add_resolved_tasks_section(document, resolved_issues_df)

                # Save document
                word_path = Path(f"{output_file}.docx")
                document.save(word_path)
                logger.info("Word report successfully created: %s", word_path)
