"""
Jira List View report - tasks grouped by assignee and week.
"""

from __future__ import annotations

from datetime import datetime, timedelta
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..export.word import add_hyperlink
from .jira_utils import norm_name, is_empty_task, get_valid_weeks


def read_member_list(member_list_file: str) -> list[str]:
    """
    Read the list of required assignees from an Excel file.
    Assumes the assignee names are in column 'E'.

    Args:
        member_list_file: Path to Excel file

    Returns:
        List of assignee names
    """
    from openpyxl import load_workbook

    wb = load_workbook(member_list_file)
    sheet = wb.active
    assignee_column = 'E'

    assignees = [
        sheet[f"{assignee_column}{row}"].value
        for row in range(2, sheet.max_row + 1)
        if sheet[f"{assignee_column}{row}"].value
    ]
    return list(set(assignees))


def add_list_view_to_document(
    document: Document,
    data: pd.DataFrame,
    start_date: str,
    end_date: str,
    jira_url: str,
    member_list_file: str | None = None,
) -> None:
    """
    Add List View section to Word document - tasks grouped by assignee and week.

    Args:
        document: Word document to add section to
        data: DataFrame with Issue_key, Summary, Assignee, Status, Week columns
        start_date: Start date string (YYYY-MM-DD)
        end_date: End date string (YYYY-MM-DD)
        jira_url: Base Jira URL for hyperlinks
        member_list_file: Optional path to Excel file with member list
    """
    document.add_heading("List View", level=2)

    # Get list of all assignees
    if member_list_file:
        required_assignees = read_member_list(member_list_file)
    else:
        required_assignees = sorted(data["Assignee"].unique())

    # Get all weeks in range
    valid_weeks = get_valid_weeks(start_date, end_date)

    for assignee in required_assignees:
        assignee_data = data[data["Assignee_norm"] == norm_name(assignee)]

        # Add assignee heading
        paragraph_assignee = document.add_paragraph()
        paragraph_assignee_format = paragraph_assignee.paragraph_format
        paragraph_assignee_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_assignee_format.space_before = Pt(12)
        paragraph_assignee_format.space_after = Pt(12)
        paragraph_assignee_format.line_spacing = 1.0
        paragraph_assignee.style = "Heading 2"

        assignee_run = paragraph_assignee.add_run(assignee)
        assignee_run.font.name = "Times New Roman"
        assignee_run.font.size = Pt(11)
        assignee_run.font.bold = True
        assignee_run.font.color.rgb = RGBColor(0, 0, 0)

        for week in valid_weeks:
            week_data = assignee_data[assignee_data["Week"] == week]
            year, week_num = map(int, week.split("-W"))
            week_start = pd.Timestamp.fromisocalendar(year, week_num, 1).strftime("%Y-%m-%d")
            week_end = (pd.Timestamp.fromisocalendar(year, week_num, 1) + timedelta(days=6)).strftime("%Y-%m-%d")
            week_header = f"ww{week_num} {week_start}-{week_end}"

            # Add week heading
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph_format.space_before = Pt(13)
            paragraph_format.space_after = Pt(13)
            paragraph_format.line_spacing = 1.73
            paragraph.style = "Heading 3"

            paragraph_run = paragraph.add_run(week_header)
            paragraph_run.font.name = "Times New Roman"
            paragraph_run.font.size = Pt(11)
            paragraph_run.font.bold = True
            paragraph_run.font.color.rgb = RGBColor(0, 0, 0)

            # Check if there are real tasks
            has_real_task = any(
                not is_empty_task(summary, status)
                for summary, status in zip(week_data["Summary"], week_data["Status"])
            )

            if week_data.empty or not has_real_task:
                # No tasks - add "vacation"
                paragraph = document.add_paragraph(style='List Bullet 2')
                vacation_run = paragraph.add_run("vacation")
                vacation_run.font.name = "Times New Roman"
                vacation_run.font.size = Pt(11)
                vacation_run.font.color.rgb = RGBColor(0, 0, 0)
                continue
            else:
                # Add tasks
                for row in week_data.itertuples(index=False, name="Row"):
                    paragraph = document.add_paragraph(style='List Bullet 2')

                    # Determine prefix based on status
                    if row.Status == "Resolved":
                        prefix = "Resolved task - "
                    elif row.Status == "In progress":
                        prefix = "Task in progress - "
                    else:
                        prefix = ""

                    # Add prefix as regular text
                    if prefix:
                        prefix_run = paragraph.add_run(prefix)
                        prefix_run.font.name = "Times New Roman"
                        prefix_run.font.size = Pt(11)

                    if isinstance(row.Issue_key, str) and row.Issue_key.strip():
                        add_hyperlink(
                            paragraph,
                            f"{jira_url}/browse/{row.Issue_key}",
                            f"{row.Issue_key} - {row.Summary}",
                            font_name="Times New Roman",
                            font_size=11
                        )

