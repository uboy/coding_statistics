"""
Jira Epic Progress report - resolved tasks grouped by epics.
"""

from __future__ import annotations

from typing import Any

import pandas as pd
from docx import Document

from ..export.word import _apply_paragraph_style


def generate_epic_resolved_hierarchy(resolved_df: pd.DataFrame) -> list[dict[str, Any]]:
    """
    Generate a summary of resolved tasks grouped by epics with parent/subtask hierarchy.

    Args:
        resolved_df: DataFrame with Issue_key, Summary, Epic_Link, Epic_Name,
            Parent_Key, Parent_Summary, Type columns

    Returns:
        List of dictionaries with Epic name and Parents list
    """
    if resolved_df.empty:
        return []

    required_cols = {"Issue_key", "Summary", "Epic_Link", "Epic_Name", "Parent_Key", "Parent_Summary", "Type"}
    if not required_cols.issubset(set(resolved_df.columns)):
        return []

    resolved_df = resolved_df.dropna(subset=["Issue_key"]).copy()
    if resolved_df.empty:
        return []

    parent_summary_map = (
        resolved_df[["Issue_key", "Summary"]]
        .dropna()
        .set_index("Issue_key")["Summary"]
        .to_dict()
    )
    epic_name_map = (
        resolved_df[["Epic_Link", "Epic_Name"]]
        .dropna()
        .drop_duplicates(subset=["Epic_Link"])
        .set_index("Epic_Link")["Epic_Name"]
        .to_dict()
    )
    parent_epic_map = (
        resolved_df[["Issue_key", "Epic_Link"]]
        .dropna()
        .set_index("Issue_key")["Epic_Link"]
        .to_dict()
    )

    epic_groups: dict[str, dict[str, Any]] = {}

    for _, row in resolved_df.iterrows():
        issue_key = row.get("Issue_key", "")
        summary = row.get("Summary", "")
        parent_key = row.get("Parent_Key", "")
        parent_summary = row.get("Parent_Summary", "")
        issue_type = row.get("Type", "")

        epic_link = row.get("Epic_Link", "")
        if not epic_link and parent_key:
            epic_link = parent_epic_map.get(parent_key, "")
        epic_name = epic_name_map.get(epic_link, "Unknown Epic")
        epic_bucket = epic_groups.setdefault(
            epic_link,
            {"Epic": epic_name, "Parents": {}}
        )

        if str(issue_type).lower() == "sub-task" and parent_key:
            parent_bucket = epic_bucket["Parents"].setdefault(
                parent_key,
                {
                    "Parent_Key": parent_key,
                    "Parent_Summary": parent_summary or parent_summary_map.get(parent_key, ""),
                    "Subtasks": [],
                },
            )
            parent_bucket["Subtasks"].append(
                {
                    "Task_Key": issue_key,
                    "Task_Summary": summary,
                }
            )
        else:
            parent_bucket = epic_bucket["Parents"].setdefault(
                issue_key,
                {
                    "Parent_Key": issue_key,
                    "Parent_Summary": summary,
                    "Subtasks": [],
                },
            )
            parent_bucket["Parent_Summary"] = summary or parent_bucket["Parent_Summary"]

    epic_summary = []
    for epic_data in epic_groups.values():
        parents_list = list(epic_data["Parents"].values())
        epic_summary.append({"Epic": epic_data["Epic"], "Parents": parents_list})

    return epic_summary


def generate_epic_progress_from_worklogs(worklogs_df: pd.DataFrame) -> list[dict[str, Any]]:
    """
    Generate a summary of in-progress tasks grouped by epics based on worklogs.

    Args:
        worklogs_df: DataFrame with Issue_key, Summary, Epic_Link, Epic_Name, Resolution columns

    Returns:
        List of dictionaries with Epic name and Tasks list
    """
    if worklogs_df.empty:
        return []

    resolution_value = worklogs_df.get("Resolution")
    if resolution_value is None:
        unresolved = worklogs_df.copy()
    else:
        unresolved = worklogs_df[resolution_value.fillna("").astype(str).eq("")]
    unresolved = unresolved.dropna(subset=["Epic_Link"])

    if unresolved.empty:
        return []

    unresolved = unresolved.drop_duplicates(subset=["Issue_key"])

    parent_summary_map = (
        unresolved[["Issue_key", "Summary"]]
        .dropna()
        .set_index("Issue_key")["Summary"]
        .to_dict()
    )
    epic_name_map = (
        unresolved[["Epic_Link", "Epic_Name"]]
        .dropna()
        .drop_duplicates(subset=["Epic_Link"])
        .set_index("Epic_Link")["Epic_Name"]
        .to_dict()
    )
    parent_epic_map = (
        unresolved[["Issue_key", "Epic_Link"]]
        .dropna()
        .set_index("Issue_key")["Epic_Link"]
        .to_dict()
    )

    epic_groups: dict[str, dict[str, Any]] = {}

    for _, row in unresolved.iterrows():
        issue_key = row["Issue_key"]
        summary = row.get("Summary", "")
        parent_key = row.get("Parent_Key", "")
        parent_summary = row.get("Parent_Summary", "")
        issue_type = row.get("Type", "")

        epic_link = row.get("Epic_Link", "")
        if not epic_link and parent_key:
            epic_link = parent_epic_map.get(parent_key, "")

        epic_name = epic_name_map.get(epic_link, "Unknown Epic")
        epic_bucket = epic_groups.setdefault(
            epic_link,
            {"Epic": epic_name, "Parents": {}}
        )

        if issue_type.lower() == "sub-task" and parent_key:
            parent_bucket = epic_bucket["Parents"].setdefault(
                parent_key,
                {
                    "Parent_Key": parent_key,
                    "Parent_Summary": parent_summary or parent_summary_map.get(parent_key, ""),
                    "Subtasks": [],
                },
            )
            parent_bucket["Subtasks"].append(
                {
                    "Task_Key": issue_key,
                    "Task_Summary": summary,
                }
            )
        else:
            parent_bucket = epic_bucket["Parents"].setdefault(
                issue_key,
                {
                    "Parent_Key": issue_key,
                    "Parent_Summary": summary,
                    "Subtasks": [],
                },
            )
            parent_bucket["Parent_Summary"] = summary or parent_bucket["Parent_Summary"]

    epic_summary = []
    for epic_data in epic_groups.values():
        parents_list = list(epic_data["Parents"].values())
        epic_summary.append({"Epic": epic_data["Epic"], "Parents": parents_list})

    return epic_summary


def add_epic_progress_to_document(
    document: Document,
    epic_summary: list[dict[str, Any]],
    jira_url: str,
    progress_summary: list[dict[str, Any]] | None = None,
) -> None:
    """
    Add Epic Progress section to Word document.

    Args:
        document: Word document to add section to
        epic_summary: List of epic dictionaries from generate_epic_resolved_hierarchy
        jira_url: Base Jira URL for hyperlinks
    """
    document.add_heading("Epic Progress", level=1)

    document.add_heading("Resolved Tasks", level=2)
    if epic_summary:
        _render_resolved_epics(document, epic_summary, heading_level=3)
    else:
        document.add_paragraph("No resolved tasks for open epics during the specified period.")

    if progress_summary is not None:
        document.add_heading("Progressed Tasks", level=2)
        if progress_summary:
            _render_progressed_epics(document, progress_summary, heading_level=3)
        else:
            document.add_paragraph("No in-progress tasks with worklogs during the specified period.")


def _render_resolved_epics(
    document: Document,
    epic_summary: list[dict[str, Any]],
    heading_level: int,
) -> None:
    for epic in epic_summary:
        document.add_heading(epic["Epic"], level=heading_level)
        for parent in epic.get("Parents", []):
            parent_paragraph = document.add_paragraph(
                f"{parent['Parent_Key']}: {parent['Parent_Summary']}",
                style="List Bullet 2",
            )
            _apply_paragraph_style([parent_paragraph], font_name="Calibri (Body)", font_size=10)

            for subtask in parent.get("Subtasks", []):
                sub_paragraph = document.add_paragraph(
                    f"{subtask['Task_Key']}: {subtask['Task_Summary']}",
                    style="List Bullet 3",
                )
                _apply_paragraph_style([sub_paragraph], font_name="Calibri (Body)", font_size=10)


def _render_progressed_epics(
    document: Document,
    epic_summary: list[dict[str, Any]],
    heading_level: int,
) -> None:
    for epic in epic_summary:
        document.add_heading(epic["Epic"], level=heading_level)
        for parent in epic.get("Parents", []):
            parent_paragraph = document.add_paragraph(
                f"{parent['Parent_Key']}: {parent['Parent_Summary']}",
                style="List Bullet 2",
            )
            _apply_paragraph_style([parent_paragraph], font_name="Calibri (Body)", font_size=10)

            for subtask in parent.get("Subtasks", []):
                sub_paragraph = document.add_paragraph(
                    f"{subtask['Task_Key']}: {subtask['Task_Summary']}",
                    style="List Bullet 3",
                )
                _apply_paragraph_style([sub_paragraph], font_name="Calibri (Body)", font_size=10)


def add_resolved_tasks_section(
    document: Document,
    resolved_tasks: pd.DataFrame,
) -> None:
    """
    Add a section to the Word document for resolved tasks, grouped by week and parent tasks.

    Args:
        document: Word document to add section to
        resolved_tasks: DataFrame with Resolution_Date, Resolution_Week, Parent_Key, Type columns
    """
    document.add_heading("Resolved Tasks", level=1)
    if resolved_tasks.empty:
        document.add_paragraph("No resolved tasks during the specified period.")
        return

    required_cols = {"Issue_key", "Summary", "Resolution_Date", "Parent_Key", "Parent_Summary", "Type"}
    if not required_cols.issubset(set(resolved_tasks.columns)):
        document.add_paragraph("No resolved tasks during the specified period.")
        return

    resolved_tasks = resolved_tasks.copy()
    if "Resolution_Week" not in resolved_tasks.columns:
        resolved_tasks.loc[:, "Resolution_Week"] = pd.to_datetime(
            resolved_tasks["Resolution_Date"], errors="coerce"
        ).dt.strftime("%G-W%V")

    from datetime import timedelta

    for week, tasks in resolved_tasks.groupby("Resolution_Week"):
        year, week_num = map(int, week.split("-W"))
        week_start = pd.Timestamp.fromisocalendar(year, week_num, 1)
        week_end = week_start + timedelta(days=6)
        week_header = f"Week {week} ({week_start.strftime('%d/%m')} - {week_end.strftime('%d/%m')})"
        document.add_heading(week_header, level=3)

        parents: dict[str, dict[str, Any]] = {}

        for _, task in tasks.iterrows():
            issue_key = task.get("Issue_key", "")
            summary = task.get("Summary", "")
            parent_key = task.get("Parent_Key", "")
            parent_summary = task.get("Parent_Summary", "")
            issue_type = task.get("Type", "")

            if str(issue_type).lower() == "sub-task" and parent_key:
                parent_bucket = parents.setdefault(
                    parent_key,
                    {
                        "Parent_Key": parent_key,
                        "Parent_Summary": parent_summary,
                        "Subtasks": [],
                    },
                )
                if not parent_bucket["Parent_Summary"]:
                    parent_bucket["Parent_Summary"] = parent_summary
                parent_bucket["Subtasks"].append(
                    {
                        "Task_Key": issue_key,
                        "Task_Summary": summary,
                    }
                )
            else:
                parent_bucket = parents.setdefault(
                    issue_key,
                    {
                        "Parent_Key": issue_key,
                        "Parent_Summary": summary,
                        "Subtasks": [],
                    },
                )
                parent_bucket["Parent_Summary"] = summary or parent_bucket["Parent_Summary"]

        for parent in parents.values():
            paragraph = document.add_paragraph(
                f"{parent['Parent_Key']}: {parent['Parent_Summary']}",
                style="List Bullet 2",
            )
            _apply_paragraph_style([paragraph], font_name="Calibri (Body)", font_size=10)

            for subtask in parent.get("Subtasks", []):
                document.add_paragraph(
                    f"{subtask['Task_Key']}: {subtask['Task_Summary']}",
                    style="List Bullet 3",
                )


