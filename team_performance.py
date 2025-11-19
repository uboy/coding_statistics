# team_performance.py

import pandas as pd
from datetime import datetime

def calculate_team_performance(data: pd.DataFrame, assignees: list[str]) -> pd.DataFrame:
    """
    Рассчитывает метрики производительности по каждому участнику команды на основе выгрузки из Jira.
    """
    # Оставляем только завершённые задачи
    resolved = data[data["Status"] == "Resolved"].copy()
    resolved["Resolution_Date"] = pd.to_datetime(resolved["Resolution_Date"])

    metrics = []

    for person in assignees:
        person_data = resolved[resolved["Assignee"] == person]

        total_tasks = len(person_data)
        bug_count = len(person_data[person_data["Type"] == "Bug"])
        reopened_count = 0  # Заменить, если появится источник переоткрытых задач
        blocked_count = len(person_data[person_data["Summary"].str.contains("block", case=False, na=False)])

        try:
            avg_resolution_days = (person_data["Resolution_Date"] - pd.to_datetime(person_data["Resolution_Date"].min())).dt.days.mean()
        except Exception:
            avg_resolution_days = None

        metrics.append({
            "Assignee": person,
            "Resolved_Tasks": total_tasks,
            "Bugs_Resolved": bug_count,
            "Blocked_Tasks": blocked_count,
            "Avg_Resolution_Time": round(avg_resolution_days, 2) if avg_resolution_days else None,
            "Reopened_Tasks": reopened_count  # Пока 0
        })

    df = pd.DataFrame(metrics)

    # Подсчёт итогового рейтинга: можно настроить веса
    df["Score"] = (
        df["Resolved_Tasks"] * 2
        - df["Bugs_Resolved"] * 1.5
        - df["Blocked_Tasks"] * 1
        - df["Reopened_Tasks"] * 2
        - df["Avg_Resolution_Time"].fillna(0) * 0.5
    )

    df["Rank"] = df["Score"].rank(method="min", ascending=False).astype(int)
    return df.sort_values("Rank")


def export_team_performance_to_excel(df: pd.DataFrame, output_file: str, sheet_name="Team Performance"):
    with pd.ExcelWriter(f"{output_file}.xlsx", engine="openpyxl", mode="a", if_sheet_exists="replace") as writer:
        df.to_excel(writer, sheet_name=sheet_name, index=False)


def add_team_performance_to_docx(df: pd.DataFrame, doc_path: str):
    from docx import Document
    from docx.shared import Pt

    doc_path = f"{doc_path}.docx"
    doc = Document(doc_path)
    doc.add_page_break()
    doc.add_heading("Team Performance Ranking", level=1)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.save(doc_path)
    print(f"Team Performance Ranking section added to: {doc_path}")
